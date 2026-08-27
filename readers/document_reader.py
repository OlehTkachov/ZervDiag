from pathlib import Path
import shutil
import subprocess
import tempfile
import unicodedata

import pymupdf
from docx import Document
from openpyxl import load_workbook


MIN_PDF_TEXT_CHARS = 80
MAX_CONTROL_NOISE_RATIO = 0.02

OCR_IMAGE_EXTENSIONS = {
    ".jpg",
    ".jpeg",
    ".png",
    ".tif",
    ".tiff",
}


class OCRRequired(Exception):
    def __init__(
        self,
        total_pages,
        preview="",
    ):
        super().__init__(
            "Document requires OCR"
        )

        self.total_pages = int(
            total_pages or 0
        )

        self.preview = (
            preview or ""
        )


def _is_disallowed_control(character):
    if character in {
        "\t",
        "\n",
        "\r",
    }:
        return False

    return (
        unicodedata.category(
            character
        ) == "Cc"
    )


def _control_noise_ratio(text):
    text = text or ""

    if not text:
        return 0.0

    bad = sum(
        1
        for character in text
        if _is_disallowed_control(
            character
        )
    )

    return (
        bad
        / max(1, len(text))
    )


def _looks_binary_like(text):
    return (
        _control_noise_ratio(text)
        > MAX_CONTROL_NOISE_RATIO
    )


def clean_extracted_text(text):
    """
    SQLite length(TEXT) effectively stops at embedded NUL characters,
    while Python len() counts them. Persisting such strings can therefore
    create a false `ok -> pending -> ok` cycle.

    Keep normal whitespace, remove embedded control characters, and leave
    printable Unicode untouched.
    """
    text = text or ""

    return "".join(
        character
        for character in text
        if not _is_disallowed_control(
            character
        )
    )


def read_pdf_text(
    filepath,
):
    """
    Читает ТОЛЬКО текстовый слой PDF.
    OCR здесь намеренно не запускается.
    """

    doc = pymupdf.open(
        filepath
    )

    try:
        total_pages = len(
            doc
        )

        text = []

        for page in doc:
            page_text = page.get_text(
                "text"
            )

            if page_text:
                text.append(
                    page_text
                )

        result = "\n".join(
            text
        )

    finally:
        doc.close()

    # Some PDFs expose a nominal "text layer" that is really binary/font
    # garbage. Do not index thousands of control bytes as valid text.
    # Such PDFs are better handled by the OCR queue.
    if _looks_binary_like(
        result
    ):
        raise OCRRequired(
            total_pages=total_pages,
            preview="",
        )

    result = clean_extracted_text(
        result
    )

    if (
        len(
            result.strip()
        )
        >= MIN_PDF_TEXT_CHARS
    ):
        return result

    raise OCRRequired(
        total_pages=total_pages,
        preview=result,
    )


def read_docx(
    filepath,
):
    text = []

    try:
        document = Document(
            filepath
        )

        for paragraph in document.paragraphs:
            if paragraph.text:
                text.append(
                    paragraph.text
                )

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        text.append(
                            cell.text
                        )

    except Exception as error:
        print(
            f"DOCX_READ_ERROR: "
            f"{filepath}: {error}",
            flush=True,
        )
        return ""

    return "\n".join(
        text
    )


def read_xlsx(
    filepath,
):
    text = []

    try:
        workbook = load_workbook(
            filepath,
            read_only=True,
            data_only=True,
        )

        for sheet in workbook.worksheets:
            text.append(
                f"\n[{sheet.title}]"
            )

            for row in sheet.iter_rows(
                values_only=True
            ):
                values = [
                    str(value)
                    for value in row
                    if value is not None
                ]

                if values:
                    text.append(
                        " | ".join(
                            values
                        )
                    )

        workbook.close()

    except Exception as error:
        print(
            f"XLSX_READ_ERROR: "
            f"{filepath}: {error}",
            flush=True,
        )
        return ""

    return "\n".join(
        text
    )


def find_libreoffice():
    soffice = shutil.which(
        "soffice"
    )

    if soffice:
        return soffice

    candidates = [
        Path(
            r"C:\Program Files\LibreOffice"
            r"\program\soffice.exe"
        ),
        Path(
            r"C:\Program Files (x86)\LibreOffice"
            r"\program\soffice.exe"
        ),
    ]

    for path in candidates:
        if path.exists():
            return str(
                path
            )

    return None


def _convert_with_libreoffice(
    filepath,
    target_format,
    target_suffix,
):
    soffice = find_libreoffice()

    if not soffice:
        print(
            "LIBREOFFICE_NOT_FOUND",
            flush=True,
        )
        return None, None

    source = Path(
        filepath
    )

    output_dir = Path(
        tempfile.mkdtemp(
            prefix="zervdiag_convert_"
        )
    )

    try:
        result = subprocess.run(
            [
                soffice,
                "--headless",
                "--convert-to",
                target_format,
                "--outdir",
                str(output_dir),
                str(source),
            ],
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            timeout=120,
        )

        if result.returncode != 0:
            print(
                f"LIBREOFFICE_ERROR: "
                f"{filepath}: {result.stderr}",
                flush=True,
            )

            shutil.rmtree(
                output_dir,
                ignore_errors=True,
            )

            return None, None

        candidates = list(
            output_dir.glob(
                f"*{target_suffix}"
            )
        )

        if not candidates:
            print(
                f"LIBREOFFICE_NO_OUTPUT: "
                f"{filepath}",
                flush=True,
            )

            shutil.rmtree(
                output_dir,
                ignore_errors=True,
            )

            return None, None

        return (
            candidates[0],
            output_dir,
        )

    except Exception as error:
        print(
            f"LIBREOFFICE_EXCEPTION: "
            f"{filepath}: {error}",
            flush=True,
        )

        shutil.rmtree(
            output_dir,
            ignore_errors=True,
        )

        return None, None


def read_doc(
    filepath,
):
    """
    Legacy .DOC is converted to .DOCX first, then parsed structurally.

    The previous DOC -> TXT path depended on LibreOffice text-export
    encoding. On old Cyrillic documents that could replace thousands of
    characters with U+FFFD and silently lose diagnostic codes. DOCX keeps
    Unicode text and tables intact for python-docx.
    """
    converted, directory = (
        _convert_with_libreoffice(
            filepath,
            "docx",
            ".docx",
        )
    )

    if not converted:
        return ""

    try:
        return read_docx(
            converted
        )

    finally:
        shutil.rmtree(
            directory,
            ignore_errors=True,
        )


def read_xls(
    filepath,
):
    converted, directory = (
        _convert_with_libreoffice(
            filepath,
            "xlsx",
            ".xlsx",
        )
    )

    if not converted:
        return ""

    try:
        return read_xlsx(
            converted
        )

    finally:
        shutil.rmtree(
            directory,
            ignore_errors=True,
        )


def read_odt(
    filepath,
):
    return read_doc(
        filepath
    )


def read_ods(
    filepath,
):
    return read_xls(
        filepath
    )


def read_text(
    filepath,
):
    path = Path(
        filepath
    )

    for encoding in (
        "utf-8-sig",
        "utf-8",
        "cp1251",
        "cp866",
        "latin-1",
    ):
        try:
            text = path.read_text(
                encoding=encoding
            )

            if _looks_binary_like(
                text
            ):
                print(
                    f"TEXT_BINARY_LIKE: "
                    f"{filepath}",
                    flush=True,
                )
                return ""

            return clean_extracted_text(
                text
            )

        except UnicodeDecodeError:
            continue

        except Exception as error:
            print(
                f"TEXT_READ_ERROR: "
                f"{filepath}: {error}",
                flush=True,
            )
            return ""

    return ""


def read_document(
    filepath,
    stop_callback=None,
):
    """
    Быстрое извлечение для ОСНОВНОЙ индексации.

    PDF без нормального текстового слоя вызывает OCRRequired.
    JPG/PNG/TIF/TIFF тоже вызывают OCRRequired.
    Tesseract здесь никогда не запускается.
    """

    extension = Path(
        filepath
    ).suffix.lower()

    if extension == ".pdf":
        return read_pdf_text(
            filepath
        )

    if extension == ".docx":
        result = read_docx(
            filepath
        )

    elif extension == ".doc":
        result = read_doc(
            filepath
        )

    elif extension == ".xlsx":
        result = read_xlsx(
            filepath
        )

    elif extension == ".xls":
        result = read_xls(
            filepath
        )

    elif extension == ".odt":
        result = read_odt(
            filepath
        )

    elif extension == ".ods":
        result = read_ods(
            filepath
        )

    elif extension in {
        ".txt",
        ".csv",
        ".json",
    }:
        result = read_text(
            filepath
        )

    elif extension in OCR_IMAGE_EXTENSIONS:
        raise OCRRequired(
            total_pages=0,
            preview="",
        )

    else:
        result = ""

    return clean_extracted_text(
        result
    )
